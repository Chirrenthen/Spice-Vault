"""
Auto DOCX Biller — Single-file Reference Implementation (ASCII-safe)
-------------------------------------------------------------------
This file contains a complete, runnable implementation for generating Word
(.docx) invoices from a .docx template with {{placeholders}}. It also includes:
- Robust placeholder replacement (paragraphs, tables, headers, footers)
- Repeating table rows using {{#items}} ... {{/items}} markers
- CLI that asks questions, computes totals, and saves a new .docx
- Built-in self tests (unittest) you can run with --self-test
- A sample fields.json you can write with --write-sample-config

IMPORTANT: All content below is ASCII-only to avoid SyntaxError from curly quotes.

How to use (quick start):
  1) Create or open your Word template (template.docx) and insert placeholders:
       Invoice No: {{invoice_no}}
       Date: {{invoice_date}}
       Bill To: {{bill_to_name}}
       Address: {{bill_to_addr}}
       Subtotal: {{subtotal}}  Tax: {{tax}}  Total: {{total}}

     For line items, make a small table like this (one row template):
       Row A: any cell contains {{#items}}
       Row B: cells use {{description}}, {{qty}}, {{rate}}, {{amount}}
       Row C: any cell contains {{/items}}

  2) Save a JSON config (fields.json). You can auto-generate a sample with:
       python this_file.py --write-sample-config fields.json

  3) Fill the document from CLI:
       python this_file.py --template template.docx --config fields.json

  4) Run self-tests anytime:
       python this_file.py --self-test

"""

from __future__ import annotations

import argparse
import json
import os
import re
import tempfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# ------------------------------
# Placeholder patterns (ASCII)
# ------------------------------
PLACEHOLDER = re.compile(r"{{\s*([a-zA-Z0-9_\.]+)\s*}}")
SECTION_START = re.compile(r"{{\s*#([a-zA-Z0-9_]+)\s*}}")
SECTION_END = re.compile(r"{{\s*/([a-zA-Z0-9_]+)\s*}}")


class DocxTemplateFiller:
    """Minimal, accurate .docx templating engine.

    Features
    - Replace {{var}} in paragraphs, table cells, headers, and footers.
    - Repeat table rows between {{#items}} and {{/items}} using the row(s) in between
      as the template for each item in context["items"].

    Limitations
    - Text boxes/shapes are not supported by python-docx. Use normal paragraphs or tables.
    - Complex per-run formatting in a single paragraph may be simplified in replaced text.
    """

    def __init__(self, template_path: str):
        self.doc = Document(template_path)

    # Public API
    def fill(self, context: Dict[str, Any]):
        # Order matters: expand repeaters first, then perform simple replacements
        self._process_tables(context)
        self._replace_in_headers_footers(context)
        self._replace_in_body(context)

    def save(self, out_path: str):
        self.doc.save(out_path)

    # Internals
    def _iter_paragraphs_in(self, container) -> List[Paragraph]:
        for p in container.paragraphs:
            yield p
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def _replace_in_headers_footers(self, ctx: Dict[str, Any]):
        for section in self.doc.sections:
            for p in self._iter_paragraphs_in(section.header):
                self._replace_in_paragraph(p, ctx)
            for p in self._iter_paragraphs_in(section.footer):
                self._replace_in_paragraph(p, ctx)

    def _replace_in_body(self, ctx: Dict[str, Any]):
        for p in self._iter_paragraphs_in(self.doc):
            self._replace_in_paragraph(p, ctx)

    def _replace_in_paragraph(self, paragraph: Paragraph, ctx: Dict[str, Any]):
        if not paragraph.runs:
            return
        text = "".join(run.text for run in paragraph.runs)
        replaced = self._replace_placeholders(text, ctx)
        if replaced != text:
            # Rebuild paragraph with a single run containing the replaced text
            # (simple and reliable; preserves paragraph-level formatting)
            for _ in range(len(paragraph.runs)):
                r = paragraph.runs[0]
                r.text = ""
                r.element.getparent().remove(r.element)
            paragraph.add_run(replaced)

    def _replace_placeholders(self, s: str, ctx: Dict[str, Any]) -> str:
        def _lookup(path: str, data: Dict[str, Any]) -> str:
            cur: Any = data
            for part in path.split("."):
                if isinstance(cur, dict):
                    cur = cur.get(part, "")
                else:
                    return ""
            return str(cur)

        return PLACEHOLDER.sub(lambda m: _lookup(m.group(1), ctx), s)

    # Table repeater logic
    def _process_tables(self, ctx: Dict[str, Any]):
        for tbl in self.doc.tables:
            self._expand_table_sections(tbl, ctx)
            # And do a pass of simple replacements inside all table cells
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph(p, ctx)

    def _expand_table_sections(self, table: Table, ctx: Dict[str, Any]):
        rows = list(table.rows)
        i = 0
        while i < len(rows):
            row = rows[i]
            row_text = self._row_text(row)
            start = SECTION_START.search(row_text)
            if not start:
                i += 1
                continue

            section_name = start.group(1)
            # Find end marker
            j = i + 1
            end_idx = None
            while j < len(rows):
                r2 = rows[j]
                r2_text = self._row_text(r2)
                end = SECTION_END.search(r2_text)
                if end and end.group(1) == section_name:
                    end_idx = j
                    break
                j += 1

            if end_idx is None:
                # No matching end; skip this marker to avoid infinite loop
                i += 1
                continue

            # Determine template rows between start and end. If none, use the start row itself.
            template_rows = rows[i + 1 : end_idx] or [rows[i]]

            # Remove section markers from template rows
            for tr in template_rows:
                for cell in tr.cells:
                    for p in cell.paragraphs:
                        self._strip_markers_in_paragraph(p)

            items = ctx.get(section_name, [])
            insert_pos = end_idx + 1
            for item in items:
                for tr in template_rows:
                    new_tr = self._clone_row(table, tr)
                    merged = {**ctx, **item}
                    for cell in new_tr.cells:
                        for p in cell.paragraphs:
                            self._replace_in_paragraph(p, merged)
                    table._tbl.insert(insert_pos, new_tr._tr)
                    insert_pos += 1

            # Remove the marker block [start..end]
            for k in range(i, end_idx + 1):
                table._tbl.remove(rows[k]._tr)

            rows = list(table.rows)
            i = insert_pos - (end_idx - i + 1)

    def _strip_markers_in_paragraph(self, paragraph: Paragraph):
        text = "".join(run.text for run in paragraph.runs)
        text = SECTION_START.sub("", text)
        text = SECTION_END.sub("", text)
        for _ in range(len(paragraph.runs)):
            r = paragraph.runs[0]
            r.text = ""
            r.element.getparent().remove(r.element)
        paragraph.add_run(text)

    def _clone_row(self, table: Table, row):
        new_tr = deepcopy(row._tr)
        return table._RowFactory(new_tr)

    def _row_text(self, row) -> str:
        parts: List[str] = []
        for cell in row.cells:
            for p in cell.paragraphs:
                parts.append("".join(r.text for r in p.runs))
        return "\n".join(parts)


# ------------------------------
# CLI utilities
# ------------------------------

def load_config(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def prompt_answers(cfg: dict) -> dict:
    answers: Dict[str, Any] = {}

    # Scalars first
    for f in cfg.get("fields", []):
        key = f["key"]
        label = f.get("label", key)
        default = f.get("default")
        hint = f" [{default}]" if default is not None else ""
        val = input(f"{label}{hint}: ").strip()
        if val == "" and default is not None:
            val = str(default)
        if val == "{today}":
            fmt = f.get("date_format", "%Y-%m-%d")
            val = datetime.now().strftime(fmt)
        answers[key] = val

    # Items (repeaters)
    items_cfg = cfg.get("items", {})
    item_fields = items_cfg.get("fields", [])
    items_key = items_cfg.get("context_key", "items")
    items: List[Dict[str, Any]] = []
    if item_fields:
        print("\nEnter line items. Leave Item Description empty to finish.")
        idx = 1
        while True:
            record: Dict[str, Any] = {}
            # Description is the sentinel for stop
            desc = input(f"Item {idx} - Item Description: ").strip()
            if desc == "":
                break
            record["description"] = desc
            for fld in item_fields:
                k = fld["key"]
                if k == "description":
                    continue
                label = fld.get("label", k)
                default = fld.get("default")
                hint = f" [{default}]" if default is not None else ""
                v = input(f"Item {idx} - {label}{hint}: ").strip()
                if v == "" and default is not None:
                    v = str(default)
                record[k] = v
            # Compute amount (qty * rate) if not provided
            try:
                qty = float(record.get("qty", 0) or 0)
                rate = float(record.get("rate", 0) or 0)
                record["amount"] = qty * rate
            except ValueError:
                record.setdefault("amount", "")
            items.append(record)
            idx += 1
    answers[items_key] = items

    # Totals
    totals_cfg = cfg.get("totals", {})
    fmt = totals_cfg.get("number_format", "{:.2f}")
    tax_rate = float(totals_cfg.get("tax_rate", 0))
    subtotal = sum(float(i.get("amount", 0) or 0) for i in items)
    tax = subtotal * tax_rate
    total = subtotal + tax
    answers.update({
        "subtotal": fmt.format(subtotal),
        "tax": fmt.format(tax),
        "total": fmt.format(total),
    })

    return answers


def run_fill(template: str, config_path: str, out_path: str | None = None):
    cfg = load_config(config_path)
    answers = prompt_answers(cfg)

    out_name_tpl = cfg.get("output_name_template") or "invoice_{invoice_no}.docx"
    out_path = out_path or out_name_tpl.format(**answers)

    filler = DocxTemplateFiller(template)
    filler.fill(answers)
    filler.save(out_path)
    print(f"\nDone. Wrote: {out_path}")


# ------------------------------
# Sample config writer (ASCII only)
# ------------------------------
SAMPLE_FIELDS_JSON = {
    "output_name_template": "invoice_{invoice_no}.docx",
    "fields": [
        {"key": "invoice_no",   "label": "Invoice Number"},
        {"key": "invoice_date", "label": "Invoice Date", "default": "{today}", "date_format": "%d-%m-%Y"},
        {"key": "bill_to_name", "label": "Bill To - Name"},
        {"key": "bill_to_addr", "label": "Bill To - Address"},
        {"key": "gstin",        "label": "GSTIN (optional)", "default": ""}
    ],
    "items": {
        "context_key": "items",
        "fields": [
            {"key": "description", "label": "Item Description"},
            {"key": "qty",         "label": "Qty",  "default": 1},
            {"key": "rate",        "label": "Rate"}
        ]
    },
    "totals": {"tax_rate": 0.18, "number_format": "{:.2f}"}
}


def write_sample_config(path: str):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(SAMPLE_FIELDS_JSON, f, indent=2)
    print(f"Wrote sample config to {path}")


# ------------------------------
# Self tests (unittest)
# ------------------------------

def _all_text(doc: Document) -> str:
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            parts.append(p.text)
        for p in sec.footer.paragraphs:
            parts.append(p.text)
    return "\n".join(parts)


def run_self_tests():
    import unittest

    class TestDocxFiller(unittest.TestCase):
        def setUp(self):
            self.tmpdir = tempfile.TemporaryDirectory()
            self.tmp = Path(self.tmpdir.name)

        def tearDown(self):
            self.tmpdir.cleanup()

        def test_paragraph_replacement(self):
            tpl = self.tmp / "tpl1.docx"
            out = self.tmp / "out1.docx"

            doc = Document()
            doc.add_paragraph("Invoice No: {{invoice_no}}")
            doc.add_paragraph("Date: {{invoice_date}}")
            doc.save(tpl)

            filler = DocxTemplateFiller(str(tpl))
            ctx = {"invoice_no": "1001", "invoice_date": "2025-08-24"}
            filler.fill(ctx)
            filler.save(str(out))

            res = Document(out)
            text = _all_text(res)
            self.assertIn("Invoice No: 1001", text)
            self.assertIn("Date: 2025-08-24", text)
            self.assertNotIn("{{invoice_no}}", text)

        def test_table_repeater(self):
            tpl = self.tmp / "tpl2.docx"
            out = self.tmp / "out2.docx"

            doc = Document()
            t = doc.add_table(rows=3, cols=4)
            # Row 0: start marker
            t.rows[0].cells[0].text = "{{#items}}"
            # Row 1: template row
            t.rows[1].cells[0].text = "{{description}}"
            t.rows[1].cells[1].text = "{{qty}}"
            t.rows[1].cells[2].text = "{{rate}}"
            t.rows[1].cells[3].text = "{{amount}}"
            # Row 2: end marker
            t.rows[2].cells[0].text = "{{/items}}"
            doc.add_paragraph("Total: {{total}}")
            doc.save(tpl)

            filler = DocxTemplateFiller(str(tpl))
            ctx = {
                "items": [
                    {"description": "Widget A", "qty": 2, "rate": 50, "amount": 100},
                    {"description": "Widget B", "qty": 1, "rate": 80, "amount": 80},
                ],
                "total": "180.00",
            }
            filler.fill(ctx)
            filler.save(str(out))

            res = Document(out)
            text = _all_text(res)
            self.assertIn("Widget A", text)
            self.assertIn("Widget B", text)
            self.assertIn("Total: 180.00", text)
            self.assertNotIn("{{#items}}", text)
            self.assertNotIn("{{/items}}", text)
            self.assertNotIn("{{description}}", text)

    suite = unittest.defaultTestLoader.loadTestsFromTestCase(TestDocxFiller)
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    if not result.wasSuccessful():
        raise SystemExit(1)


# ------------------------------
# Main entry
# ------------------------------

def main():
    parser = argparse.ArgumentParser(description="Auto DOCX Biller - ask, fill, export .docx (ASCII-safe)")
    parser.add_argument("--template", help="Path to your .docx template with placeholders")
    parser.add_argument("--config", default="fields.json", help="Path to fields.json mapping")
    parser.add_argument("--out", default=None, help="Output .docx path (optional)")
    parser.add_argument("--write-sample-config", dest="sample_cfg", help="Write a sample fields.json to this path and exit")
    parser.add_argument("--self-test", action="store_true", help="Run built-in unit tests and exit")
    args = parser.parse_args()

    if args.sample_cfg:
        write_sample_config(args.sample_cfg)
        return

    if args.self_test:
        run_self_tests()
        return

    if not args.template:
        parser.error("--template is required unless using --self-test or --write-sample-config")

    run_fill(args.template, args.config, args.out)


if __name__ == "__main__":
    main()
